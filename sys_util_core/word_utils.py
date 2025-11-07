"""
MS Word Document Utilities
MS Word 문서 유틸리티

This module provides utility functions for creating and manipulating MS Word documents.
MS Word 문서를 생성하고 조작하기 위한 유틸리티 함수들을 제공합니다.

Note: Requires python-docx package for full functionality
참고: 전체 기능을 위해 python-docx 패키지가 필요합니다
"""

import os
from typing import List, Optional, Dict, Any


def create_word_from_array(text_array: List[str], output_file: str,
                           font_name: str = 'Arial', font_size: int = 12) -> bool:
    '''
    Create Word document from array of strings. 문자열 배열에서 Word 문서를 생성합니다.
    Args:
    text_array: Array of text strings (each becomes a paragraph) 텍스트 문자열 배열 (각각 단락이 됨)
    output_file: Output Word file path 출력 Word 파일 경로
    font_name: Font name 폰트 이름
    font_size: Font size 폰트 크기
    Returns:
    True if successful, False otherwise 성공하면 True, 실패하면 False
    '''
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        
        for text in text_array:
            paragraph = doc.add_paragraph(text)
            # Set font properties
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
        
        doc.save(output_file)
        return True
    except ImportError:
        # Fallback: create basic text file
        try:
            with open(output_file.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
                f.write('\n\n'.join(text_array))
            return True
        except Exception:
            return False
    except Exception:
        return False


def add_heading_to_word(file_path: str, heading_text: str, level: int = 1) -> bool:
    '''
    Add a heading to existing Word document. 기존 Word 문서에 제목을 추가합니다.
    Args:
    file_path: Path to Word document Word 문서 경로
    heading_text: Heading text 제목 텍스트
    level: Heading level (1-9) 제목 레벨 (1-9)
    Returns:
    True if successful, False otherwise 성공하면 True, 실패하면 False
    '''
    try:
        from docx import Document
        
        doc = Document(file_path) if os.path.exists(file_path) else Document()
        doc.add_heading(heading_text, level=level)
        doc.save(file_path)
        return True
    except Exception:
        return False


def add_table_to_word(file_path: str, data: List[List[str]], 
                      has_header: bool = True) -> bool:
    '''
    Add a table to Word document. Word 문서에 표를 추가합니다.
    Args:
    file_path: Path to Word document Word 문서 경로
    data: 2D array of table data 표 데이터의 2D 배열
    has_header: Whether first row is header 첫 번째 행이 헤더인지 여부
    Returns:
    True if successful, False otherwise 성공하면 True, 실패하면 False
    '''
    try:
        from docx import Document
        
        doc = Document(file_path) if os.path.exists(file_path) else Document()
        
        if not data:
            return False
        
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Light Grid Accent 1'
        
        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = str(cell_data)
        
        doc.save(file_path)
        return True
    except Exception:
        return False


def add_image_to_word(file_path: str, image_path: str, 
                      width_inches: Optional[float] = None) -> bool:
    '''
    Add an image to Word document. Word 문서에 이미지를 추가합니다.
    Args:
    file_path: Path to Word document Word 문서 경로
    image_path: Path to image file 이미지 파일 경로
    width_inches: Image width in inches (None for original size) 인치 단위 이미지 너비 (None이면 원본 크기)
    Returns:
    True if successful, False otherwise 성공하면 True, 실패하면 False
    '''
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document(file_path) if os.path.exists(file_path) else Document()
        
        if width_inches:
            doc.add_picture(image_path, width=Inches(width_inches))
        else:
            doc.add_picture(image_path)
        
        doc.save(file_path)
        return True
    except Exception:
        return False


def read_word_document(file_path: str) -> Optional[List[str]]:
    '''
    Read all paragraphs from Word document. Word 문서에서 모든 단락을 읽습니다.
    Args:
    file_path: Path to Word document Word 문서 경로
    Returns:
    List of paragraph texts or None if error 단락 텍스트 리스트, 에러시 None
    '''
    try:
        from docx import Document
        
        doc = Document(file_path)
        return [para.text for para in doc.paragraphs]
    except Exception:
        return None


def replace_text_in_word(file_path: str, old_text: str, new_text: str) -> bool:
    '''
    Find and replace text in Word document. Word 문서에서 텍스트를 찾아 바꿉니다.
    Args:
    file_path: Path to Word document Word 문서 경로
    old_text: Text to find 찾을 텍스트
    new_text: Replacement text 교체할 텍스트
    Returns:
    True if successful, False otherwise 성공하면 True, 실패하면 False
    '''
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        for para in doc.paragraphs:
            if old_text in para.text:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
        
        doc.save(file_path)
        return True
    except Exception:
        return False


def merge_word_documents(input_files: List[str], output_file: str) -> bool:
    '''
    Merge multiple Word documents into one. 여러 Word 문서를 하나로 병합합니다.
    Args:
    input_files: List of Word document paths Word 문서 경로 리스트
    output_file: Output Word file path 출력 Word 파일 경로
    Returns:
    True if successful, False otherwise 성공하면 True, 실패하면 False
    '''
    try:
        from docx import Document
        from docx.oxml.xmlchemy import OxmlElement
        
        merged_doc = Document()
        
        for file_path in input_files:
            sub_doc = Document(file_path)
            
            for element in sub_doc.element.body:
                merged_doc.element.body.append(element)
        
        merged_doc.save(output_file)
        return True
    except Exception:
        return False


def get_word_document_info(file_path: str) -> Optional[Dict[str, Any]]:
    '''
    Get information about Word document. Word 문서에 대한 정보를 가져옵니다.
    Args:
    file_path: Path to Word document Word 문서 경로
    Returns:
    Dictionary with document information or None if error 문서 정보 딕셔너리, 에러시 None
    '''
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        info = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'file_size_kb': os.path.getsize(file_path) / 1024 if os.path.exists(file_path) else 0
        }
        
        return info
    except Exception:
        return None


def create_word_template(output_file: str, placeholders: List[str]) -> bool:
    '''
    Create a Word document template with placeholders. 자리 표시자가 있는 Word 문서 템플릿을 생성합니다.
    Args:
    output_file: Output Word file path 출력 Word 파일 경로
    placeholders: List of placeholder names 자리 표시자 이름 리스트
    Returns:
    True if successful, False otherwise 성공하면 True, 실패하면 False
    '''
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('Document Template', 0)
        
        for placeholder in placeholders:
            doc.add_paragraph(f'{{{{{placeholder}}}}}')
        
        doc.save(output_file)
        return True
    except Exception:
        return False
